from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import pandas as pd

# Пътища
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
BASE_DIR = os.path.dirname(SCRIPT_DIR)
REPORT_DIR = os.path.join(BASE_DIR, 'report')
DATA_DIR = os.path.join(BASE_DIR, 'data')

def set_academic_style(doc):
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

def add_interactive_toc(doc):
    doc.add_heading('СЪДЪРЖАНИЕ', level=1)
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    doc.add_paragraph('Забележка: Отворете файла в Word, кликнете с десен бутон върху съдържанието и изберете "Update Field" (Обнови полето), за да се генерира автоматично.')

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run()
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)

def add_data_sample(doc, file_name, title):
    file_path = os.path.join(DATA_DIR, file_name)
    if os.path.exists(file_path):
        df = pd.read_csv(file_path).head(10)
        doc.add_heading(f'Извадка от данни: {title}', level=3)
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, column in enumerate(df.columns):
            hdr_cells[i].text = column
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(round(value, 2)) if isinstance(value, float) else str(value)
        doc.add_paragraph('')

def create_report():
    doc = Document()
    set_academic_style(doc)
    
    # СТРАНИЦА 1: ЗАГЛАВНА
    doc.add_paragraph('\n' * 5)
    title = doc.add_heading('РЕФЕРАТ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('\n' * 2)
    subtitle = doc.add_paragraph('Тема: Пълен академичен анализ на Дървовидните Алгоритми в Машинното Обучение')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(22)
    doc.add_paragraph('\n' * 2)
    course = doc.add_paragraph('Курс: Основи на изкуствения интелект')
    course.alignment = WD_ALIGN_PARAGRAPH.CENTER
    course.runs[0].font.size = Pt(14)
    doc.add_paragraph('\n' * 8)
    doc.add_paragraph('Дата: сряда, 6 май 2026 г.').alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # СТРАНИЦА 2: СЪДЪРЖАНИЕ
    add_interactive_toc(doc)
    doc.add_page_break()

    # --- 1. Въведение ---
    doc.add_heading('1. Въведение и Исторически контекст', level=1)
    doc.add_paragraph(
        'Дърветата на решенията (Decision Trees) представляват един от най-старите и същевременно най-издръжливите методи в областта на изкуствения интелект. '
        'Тяхното развитие започва още през 60-те години на миналия век с появата на алгоритми като AID (Automatic Interaction Detector) и по-късно се формализира '
        'през 80-те години от Лео Брейман и неговите колеги чрез концепцията CART (Classification and Regression Trees). '
        'Днес дървовидните модели са гръбнакът на модерните AI системи за структурирани данни, конкурирайки се успешно дори с невронните мрежи.'
    )
    doc.add_paragraph(
        'Основната парадигма на дърветата е "Разделяй и владей" (Divide and Conquer). Алгоритъмът систематично разчленява пространството от признаци на хипер-правоъгълници, '
        'във всеки от които предсказанието е константа. Този геометричен подход позволява на модела да улавя сложни, нелинейни зависимости, които са недостъпни за линейните модели.'
    )

    # --- 2. Класически алгоритми ---
    doc.add_heading('2. Класически алгоритми: ID3, C4.5 и CART', level=1)
    doc.add_paragraph('В развитието на дърветата на решенията се открояват три основни софтуерни и математически архитектури:')
    p = doc.add_paragraph()
    p.add_run('ID3 (Iterative Dichotomiser 3): ').bold = True
    p.add_run('Създаден от Рос Куинлан, този алгоритъм използва Информационната печалба (Information Gain) като основен критерий. '
              'Основният му недостатък е, че работи само с номинални (категорийни) признаци и е склонен към предпочитане на признаци с много категории.')
    p = doc.add_paragraph()
    p.add_run('C4.5: ').bold = True
    p.add_run('Еволюцията на ID3, която въвежда работа с числови стойности (чрез прагови разделяния) и механизъм за подрязване (Pruning). '
              'Вместо чист Information Gain, той използва Gain Ratio, за да нормализира пристрастието към много категории.')
    p = doc.add_paragraph()
    p.add_run('CART: ').bold = True
    p.add_run('Това е архитектурата, използвана в днешния софтуер (като scikit-learn). CART генерира стриктно бинарни дървета (всеки възел има точно два клона) '
              'и поддържа както класификация (чрез Gini), така и регресия (чрез MSE).')

    # --- 3. Математически фундамент ---
    doc.add_heading('3. Математически фундамент и оптимизационни критерии', level=1)
    doc.add_paragraph('Процесът на обучение е рекурсивно разделяне. Алгоритъмът търси оптималния сплит (S), който минимизира загубата на информация.')
    
    doc.add_heading('3.1. Критерии при Класификация', level=2)
    doc.add_paragraph('При класификацията дървото се стреми към хомогенност. Използват се следните метрики:')
    p = doc.add_paragraph()
    p.add_run('Gini Impurity (Нечистота на Джини): ').bold = True
    p.add_run('Измерва вероятността случаен елемент да бъде грешно класифициран. Стойност 0 означава перфектна чистота. Формула: Gini(D) = 1 - Σ (p_i)^2. '
              'Тя е изключително ефективна за големи набори данни, тъй като изисква по-малко изчислителни ресурси от логаритмичните функции.')
    p = doc.add_paragraph()
    p.add_run('Entropy (Ентропия): ').bold = True
    p.add_run('Концепция от теорията на информацията на Клод Шанън. Измерва хаоса. Информационната печалба е намалението на ентропията след разделяне. Формула: Entropy = -Σ p_i * log2(p_i).')

    doc.add_heading('3.2. Критерии при Регресия', level=2)
    doc.add_paragraph('Тук целевата променлива е непрекъсната. Листата предсказват средната стойност (mean).')
    p = doc.add_paragraph()
    p.add_run('MSE (Mean Squared Error): ').bold = True
    p.add_run('Алгоритъмът минимизира сумата от квадратите на разликите между реалните стойности и средната стойност в нода.')
    doc.add_paragraph('MSE = (1/n) * Σ (y_i - ȳ)^2', style='Normal').paragraph_format.left_indent = Inches(0.5)
    p = doc.add_paragraph()
    p.add_run('MAE (Mean Absolute Error): ').bold = True
    p.add_run('По-устойчива метрика при наличие на екстремни стойности (outliers). Тя минимизира сумата от абсолютните разлики.')

    # --- 4. Overfitting ---
    doc.add_heading('4. Проблемът с преобучаването и техники за регуларизация', level=1)
    doc.add_paragraph('Дърветата на решенията са склонни да растат безкрайно, докато не обхванат всеки детайл (и шум) в данните. Това води до преобучаване (overfitting).')
    doc.add_heading('4.1. Пре-подрязване (Pre-pruning)', level=2)
    doc.add_paragraph('Задаваме лимити още по време на обучение:')
    doc.add_paragraph('• Max Depth: Ограничаваме броя на нивата (например до 5).', style='List Bullet')
    doc.add_paragraph('• Min Samples Split: Минимален брой примери за ново разделяне.', style='List Bullet')
    doc.add_paragraph('• Min Samples Leaf: Минимален брой примери в едно листо.', style='List Bullet')
    doc.add_heading('4.2. Пост-подрязване (Post-pruning)', level=2)
    doc.add_paragraph('Позволяваме на дървото да израсне напълно, след което систематично премахваме клони, които не допринасят значително за точността. Cost Complexity Pruning е типичен пример.')

    # --- 5. Ансамбли ---
    doc.add_heading('5. Ансамблово обучение: Пътят към върховите постижения', level=1)
    doc.add_paragraph('Ансамблите комбинират много "слаби" дървета, за да създадат един "силен" модел.')
    doc.add_heading('5.1. Bagging (Bootstrap Aggregating) и Random Forest', level=2)
    doc.add_paragraph('Random Forest работи чрез паралелизъм. Всяко дърво се обучава върху случайна извадка от данните и използва случайно подмножество от признаците. Когато едно дърво сгреши, останалите го коригират чрез гласуване.')
    doc.add_heading('5.2. Boosting и Gradient Boosting (XGBoost)', level=2)
    doc.add_paragraph('Boosting работи чрез последователност. Всяко ново дърво се фокусира върху примерите, които предходните са сгрешили. Градиентният бустинг (напр. XGBoost) използва оптимизация (Gradient Descent) върху загубната функция.')

    # --- 6. Приложения ---
    doc.add_heading('6. Приложения в реалния свят и индустриални гиганти', level=1)
    doc.add_paragraph('Дърветата на решенията движат някои от най-големите платформи в света:')
    p = doc.add_paragraph()
    p.add_run('Netflix и Amazon (Препоръчителни системи): ').bold = True
    p.add_run('Градиентният бустинг (XGBoost/CatBoost) често управлява финалното класиране на продуктите заради способността му да обработва милиони категорийни признаци.')
    p = doc.add_paragraph()
    p.add_run('PayPal и Банковия сектор (Откриване на измами): ').bold = True
    p.add_run('Random Forest се използва за засичане на подозрителни шаблони (fraud detection) в реално време, комбинирайки скорост с висока прецизност.')
    p = doc.add_paragraph()
    p.add_run('NASA (Анализ на космически данни): ').bold = True
    p.add_run('Класификация на небесни тела от сателитни снимки. Интерпретируемостта позволява на астрономите да разберат физическите закони.')
    p = doc.add_paragraph()
    p.add_run('Uber (Предсказване на време за пристигане - ETA): ').bold = True
    p.add_run('Огромни ансамбли от дървета се използват за предвиждане на трафика.')

    # --- 7. Практическа част ---
    doc.add_heading('7. Практическа реализация и анализ на задачите', level=1)
    
    doc.add_heading('7.1. Одобрение на кредит (Класификация чрез Decision Tree)', level=2)
    doc.add_paragraph('Целта е автоматизирана банкова система. Изборът на просто дърво е стратегически – финансовите институции са длъжни по закон да обясняват решенията си.')
    add_data_sample(doc, 'credit_approval.csv', 'Клиентски профили')
    tree_path = os.path.join(REPORT_DIR, 'credit_tree.png')
    if os.path.exists(tree_path):
        doc.add_picture(tree_path, width=Inches(6.0))
        doc.add_paragraph('Фиг. 1. Логическа структура на кредитното решение.')

    doc.add_heading('7.2. HR Анализ: Напускане на служители (Random Forest)', level=2)
    doc.add_paragraph('Използваме Random Forest, за да уловим фините нюанси между удовлетвореността и претоварването. Този модел е много по-устойчив на малки промени в субективните анкети на служителите.')
    add_data_sample(doc, 'employee_churn.csv', 'HR Статистика')
    rf_path = os.path.join(REPORT_DIR, 'rf_last_tree.png')
    if os.path.exists(rf_path):
        doc.add_picture(rf_path, width=Inches(6.0))
        doc.add_paragraph('Фиг. 2. Визуализация на един от компонентите в HR гората.')

    doc.add_heading('7.3. Оценка на имоти (XGBoost Regression)', level=2)
    doc.add_paragraph('Ценообразуването изисква висока математическа точност. XGBoost успява да намали грешката до минимум чрез градиентна оптимизация.')
    add_data_sample(doc, 'apartments.csv', 'Ценови данни')
    xgb_path = os.path.join(REPORT_DIR, 'xgboost_random_tree.png')
    if os.path.exists(xgb_path):
        doc.add_picture(xgb_path, width=Inches(6.0))
        doc.add_paragraph('Фиг. 3. Вътрешна логика на бустинг модела при определяне на цена.')

    # --- 8. Conclusion ---
    doc.add_heading('8. Заключение', level=1)
    doc.add_paragraph(
        'В хода на този анализ разгледахме пълния спектър на дървовидните алгоритми. От техните скромни начала като интерпретируеми схеми до '
        'мощните ансамбли, които доминират днешната AI индустрия. Дърветата на решенията остават един от най-важните инструменти в арсенала на всеки специалист по данни, предлагайки уникален баланс '
        'между прозрачност, скорост и изчислителна мощ.'
    )
    
    doc.add_page_break()

    # --- 9. Bibliography ---
    doc.add_heading('9. Използвана литература', level=1)
    doc.add_paragraph('1. Breiman, L., Friedman, J., Stone, C. J., & Olshen, R. A. (1984). Classification and Regression Trees. Taylor & Francis.', style='List Bullet')
    doc.add_paragraph('2. Quinlan, J. R. (1993). C4.5: Programs for Machine Learning. Morgan Kaufmann Publishers.', style='List Bullet')
    doc.add_paragraph('3. Hastie, T., Tibshirani, R., & Friedman, J. (2009). The Elements of Statistical Learning: Data Mining, Inference, and Prediction. Springer.', style='List Bullet')
    doc.add_paragraph('4. Pedregosa, F. et al. (2011). Scikit-learn: Machine Learning in Python. Journal of Machine Learning Research, 12, 2825-2830.', style='List Bullet')
    doc.add_paragraph('5. Chen, T., & Guestrin, C. (2016). XGBoost: A Scalable Tree Boosting System. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining.', style='List Bullet')
    doc.add_paragraph('6. NASA Jet Propulsion Laboratory. (2023). Machine Learning Applications in Space Exploration.', style='List Bullet')

    # Apply page numbers
    add_page_numbers(doc)
    doc.sections[0].different_first_page_header_footer = True
    
    save_path = os.path.join(REPORT_DIR, 'Decision_Trees_Professional_V3.docx')
    doc.save(save_path)
    print(f"Финалният доклад с пълно съдържание е генериран: {save_path}")

if __name__ == '__main__':
    create_report()
